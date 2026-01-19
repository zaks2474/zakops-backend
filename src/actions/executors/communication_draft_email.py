from __future__ import annotations

import asyncio
import json
import os
import threading
from pathlib import Path
from typing import Any, Dict, Optional

from actions.engine.models import ActionError, ActionPayload, ArtifactMetadata, now_utc_iso
from actions.executors._artifacts import resolve_action_artifact_dir
from actions.executors.base import ActionExecutionError, ActionExecutor, ExecutionContext, ExecutionResult
from zakops_secret_scan import find_secrets_in_text


def _run_coro_blocking(coro):
    """
    Run a coroutine from sync code.

    - If no loop is running: uses asyncio.run()
    - If a loop is already running (e.g., unit tests): runs in a dedicated thread
    """
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro)

    result: Dict[str, Any] = {}
    error: Dict[str, BaseException] = {}

    def _thread_main() -> None:
        try:
            result["value"] = asyncio.run(coro)
        except BaseException as e:  # pragma: no cover
            error["exc"] = e

    t = threading.Thread(target=_thread_main, name="draft-email-loop", daemon=True)
    t.start()
    t.join()
    if error:
        raise error["exc"]
    return result.get("value")


def _extract_json(text: str) -> Dict[str, Any]:
    raw = (text or "").strip()
    if not raw:
        return {}

    # Best-effort strip of ```json fences.
    if raw.startswith("```"):
        lines = raw.splitlines()
        if len(lines) >= 3 and lines[0].startswith("```") and lines[-1].startswith("```"):
            raw = "\n".join(lines[1:-1]).strip()
        raw = raw.lstrip("`").strip()

    try:
        data = json.loads(raw)
    except Exception:
        return {}

    return data if isinstance(data, dict) else {}


class DraftEmailExecutor(ActionExecutor):
    action_type = "COMMUNICATION.DRAFT_EMAIL"

    def validate(self, payload: ActionPayload) -> tuple[bool, Optional[str]]:
        required = ["to", "subject", "context"]
        missing = [k for k in required if not str((payload.inputs or {}).get(k, "")).strip()]
        if missing:
            return False, f"Missing required fields: {', '.join(missing)}"
        return True, None

    def execute(self, payload: ActionPayload, ctx: ExecutionContext) -> ExecutionResult:
        inputs = payload.inputs or {}
        to = str(inputs.get("to", "")).strip()
        subject = str(inputs.get("subject", "")).strip()
        context_text = str(inputs.get("context", "")).strip()
        tone = str(inputs.get("tone", "professional")).strip() or "professional"

        provider_env = (os.getenv("ZAKOPS_DRAFT_EMAIL_PROVIDER", "gemini") or "gemini").strip().lower()
        cloud_allowed = bool(getattr(ctx, "cloud_allowed", False))

        if provider_env != "stub" and not cloud_allowed:
            raise ActionExecutionError(
                ActionError(
                    code="cloud_disabled",
                    message="Cloud drafting disabled for this action (requires explicit approval of a cloud-required capability).",
                    category="cloud_policy",
                    retryable=False,
                    details={"provider": provider_env},
                )
            )

        # Secret-scan gate BEFORE any cloud send (or even stub, for safety).
        prompt = f"TO: {to}\nSUBJECT: {subject}\nTONE: {tone}\nCONTEXT:\n{context_text}\n"
        secrets = find_secrets_in_text(prompt)
        if secrets:
            raise ActionExecutionError(
                ActionError(
                    code="secret_scan_blocked",
                    message=f"Secret-like patterns detected ({', '.join(secrets)}); refusing to draft.",
                    category="cloud_policy",
                    retryable=False,
                )
            )

        cc: list[str] = []
        bcc: list[str] = []
        provider_used = "stub"
        model = "stub"
        latency_ms = 0

        if provider_env == "stub":
            body = (
                f"Hi,\n\n{context_text}\n\n"
                "Thanks,\n"
                f"{os.getenv('ZAKOPS_OPERATOR_NAME', 'ZakOps')}\n"
            )
        else:
            # Gemini provider (strict JSON).
            from chat_llm_provider import GeminiProProvider, ProviderResponse

            gemini = GeminiProProvider()
            if not getattr(gemini, "available", False):
                raise ActionExecutionError(
                    ActionError(
                        code="gemini_api_key_missing",
                        message="Gemini API key not configured (set GEMINI_API_KEY or ~/.gemini_api)",
                        category="dependency",
                        retryable=False,
                    )
                )

            system = (
                "You draft broker-facing emails.\n"
                "Return ONLY strict JSON (no markdown) matching:\n"
                '{ "subject": string, "body": string, "cc"?: string[], "bcc"?: string[] }\n'
                "Rules:\n"
                "- Keep it concise, professional, and specific.\n"
                "- Do not invent deal facts.\n"
                "- Do not include secrets.\n"
            )
            user = (
                f"To: {to}\n"
                f"Subject: {subject}\n"
                f"Tone: {tone}\n\n"
                "Context:\n"
                f"{context_text}\n"
            )

            resp = _run_coro_blocking(
                gemini.generate(
                    messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                    temperature=0.2,
                    max_tokens=2048,
                )
            )
            if not isinstance(resp, ProviderResponse):
                raise ActionExecutionError(
                    ActionError(
                        code="gemini_invalid_response",
                        message=f"Gemini returned invalid response type: {type(resp)}",
                        category="dependency",
                        retryable=False,
                    )
                )
            provider_used = str(resp.provider or "").strip() or "gemini-pro"
            model = str(resp.model or "").strip() or "gemini"
            latency_ms = int(resp.latency_ms or 0)

            data = _extract_json(resp.content)
            body = str(data.get("body") or "").strip()
            out_subject = str(data.get("subject") or "").strip()
            if out_subject:
                subject = out_subject

            cc_raw = data.get("cc")
            if isinstance(cc_raw, list):
                cc = [str(x or "").strip() for x in cc_raw if str(x or "").strip()]
            bcc_raw = data.get("bcc")
            if isinstance(bcc_raw, list):
                bcc = [str(x or "").strip() for x in bcc_raw if str(x or "").strip()]

            if not body:
                raise ActionExecutionError(
                    ActionError(
                        code="gemini_empty_draft",
                        message="Gemini returned empty draft content",
                        category="cloud_transient",
                        retryable=True,
                        details={"model": model},
                    )
                )

        out_dir = resolve_action_artifact_dir(ctx)

        md_path = out_dir / "draft_email.md"
        md_content = f"To: {to}\nSubject: {subject}\n\n{body}\n"
        md_path.write_text(md_content, encoding="utf-8")

        json_path = out_dir / "draft_email.json"
        json_path.write_text(
            json.dumps(
                {
                    "to": to,
                    "subject": subject,
                    "body": body,
                    "cc": cc,
                    "bcc": bcc,
                    "provider": provider_used,
                    "provider_env": provider_env,
                    "model": model,
                    "latency_ms": latency_ms,
                },
                indent=2,
                sort_keys=True,
            ),
            encoding="utf-8",
        )

        artifacts = [
            ArtifactMetadata(
                filename=md_path.name,
                mime_type="text/markdown",
                path=str(md_path),
                created_at=now_utc_iso(),
            )
            ,
            ArtifactMetadata(
                filename=json_path.name,
                mime_type="application/json",
                path=str(json_path),
                created_at=now_utc_iso(),
            ),
        ]

        # Optional DOCX (best-effort)
        try:
            from docx import Document  # type: ignore

            doc = Document()
            doc.add_paragraph(f"To: {to}")
            doc.add_paragraph(f"Subject: {subject}")
            doc.add_paragraph("")
            for line in body.splitlines():
                doc.add_paragraph(line)
            docx_path = out_dir / "draft_email.docx"
            doc.save(str(docx_path))
            artifacts.append(
                ArtifactMetadata(
                    filename=docx_path.name,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    path=str(docx_path),
                    created_at=now_utc_iso(),
                )
            )
        except Exception:
            pass

        outputs: Dict[str, Any] = {
            "to": to,
            "subject": subject,
            "cc": cc,
            "bcc": bcc,
            "tone": tone,
            "provider": provider_used,
            "provider_env": provider_env,
            "model": model,
            "latency_ms": latency_ms,
        }
        return ExecutionResult(outputs=outputs, artifacts=artifacts)
